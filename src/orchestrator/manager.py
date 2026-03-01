"""
Orchestrator Manager -- Edu Nexus
==================================
Connects BM25 keyword retrieval and Neo4j Graph Engine,
decides which engine to invoke, fuses their context,
and generates a real answer via Groq LLM.
"""

import asyncio
import os
import shutil
import importlib.util
import logging
from pathlib import Path
from typing import Dict, List

from dotenv import load_dotenv
from groq import Groq

from src.retrieval.bm25_index import KeywordEngine
from src.graph_engine.neo4j_ops import Neo4jConnector

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger("Orchestrator")

# ---- Project paths ----
RAW_DIR = Path("data/raw")
PROCESSED_DIR = Path("data/processed")
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

# ---- LLM config ----
LLM_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"
BM25_TOP_K = 5
GRAPH_RESULT_LIMIT = 10

SYSTEM_PROMPT = (
    "You are **Edu Nexus**, an intelligent academic assistant.\n"
    "You are provided with two types of context to answer the user's question:\n\n"
    "1. **Document Chunks** -- passages retrieved via BM25 keyword search from "
    "the course material corpus.\n"
    "2. **Knowledge-Graph Facts** -- structured triples (Subject -> Relation -> Object) "
    "extracted from a Neo4j knowledge graph.\n\n"
    "Guidelines:\n"
    "- Synthesise both contexts into a clear, concise, and accurate answer.\n"
    "- If one context is empty or irrelevant, rely on the other.\n"
    "- If neither context is sufficient, say: "
    "\"I don't have enough information in my knowledge base to answer this.\"\n"
    "- Be helpful, detailed, and academic in tone.\n"
    "- Format your answers in Markdown for readability.\n"
)


class OrchestratorManager:
    """
    Real Orchestrator that ties BM25 + Neo4j Graph Engine + Groq LLM.
    The orchestrator decides which retrieval engines to use based on
    their availability and the query, then fuses the results.
    """

    def __init__(self):
        RAW_DIR.mkdir(parents=True, exist_ok=True)
        PROCESSED_DIR.mkdir(parents=True, exist_ok=True)

        # ---- BM25 keyword engine ----
        self.keyword_engine = KeywordEngine()
        try:
            self.keyword_engine.load_index()
            self._bm25_ready = True
            logger.info("BM25 index loaded successfully.")
        except FileNotFoundError:
            self._bm25_ready = False
            logger.warning("BM25 index not found -- will be available after file upload.")

        # ---- Neo4j graph connector ----
        self.neo4j = Neo4jConnector()
        self._graph_ready = self.neo4j.verify_connectivity()
        if self._graph_ready:
            logger.info("Neo4j connection verified.")
        else:
            logger.warning("Neo4j not reachable -- graph retrieval disabled.")

        # ---- Groq LLM client ----
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise ValueError("GROQ_API_KEY not set in environment.")
        self.llm = Groq(api_key=api_key)

    # ================================================================== #
    #  RETRIEVAL                                                          #
    # ================================================================== #

    def _retrieve_bm25(self, query: str) -> List[str]:
        """Top-k document chunks from BM25 keyword search."""
        if not self._bm25_ready:
            return []
        try:
            results = self.keyword_engine.search(query, k=BM25_TOP_K)
            logger.info(f"BM25 returned {len(results)} chunks.")
            return results
        except Exception as e:
            logger.error(f"BM25 search failed: {e}")
            return []

    def _retrieve_graph(self, query: str) -> List[Dict]:
        """
        Query Neo4j for triples whose node names CONTAIN any keyword
        from the user query.
        """
        if not self._graph_ready:
            return []

        stopwords = {
            "what", "is", "the", "a", "an", "of", "to", "and", "in",
            "for", "on", "how", "does", "do", "are", "was", "were",
            "who", "which", "can", "about", "tell", "me", "explain",
            "this", "that", "it", "they", "them", "their", "its",
        }
        keywords = [
            w for w in query.split()
            if w.lower() not in stopwords and len(w) > 1
        ]
        if not keywords:
            return []

        where_clauses = " OR ".join(
            f"toLower(a.name) CONTAINS toLower($kw{i}) OR "
            f"toLower(b.name) CONTAINS toLower($kw{i})"
            for i in range(len(keywords))
        )

        cypher = (
            f"MATCH (a)-[r]->(b) "
            f"WHERE {where_clauses} "
            f"RETURN a.name AS source, type(r) AS relation, b.name AS target "
            f"LIMIT {GRAPH_RESULT_LIMIT}"
        )
        params = {f"kw{i}": kw for i, kw in enumerate(keywords)}

        try:
            results = self.neo4j.run_cypher(cypher, params)
            logger.info(f"Graph returned {len(results)} triples.")
            return results
        except Exception as e:
            logger.error(f"Graph query failed: {e}")
            return []

    # ================================================================== #
    #  ORCHESTRATION LOGIC                                                #
    # ================================================================== #

    def _decide_strategy(self, query: str) -> str:
        """
        Decide which retrieval engines to invoke based on availability.
        Returns one of: 'both', 'bm25_only', 'graph_only', 'none'.
        """
        if self._bm25_ready and self._graph_ready:
            return "both"
        elif self._bm25_ready:
            return "bm25_only"
        elif self._graph_ready:
            return "graph_only"
        else:
            return "none"

    @staticmethod
    def _format_bm25_context(chunks: List[str]) -> str:
        if not chunks:
            return "_No document chunks retrieved._"
        lines = []
        for i, chunk in enumerate(chunks, 1):
            lines.append(f"**[Chunk {i}]**\n{chunk}")
        return "\n\n".join(lines)

    @staticmethod
    def _format_graph_context(triples: List[Dict]) -> str:
        if not triples:
            return "_No knowledge-graph facts retrieved._"
        lines = []
        for t in triples:
            src = t.get("source", "?")
            rel = t.get("relation", "?")
            tgt = t.get("target", "?")
            lines.append(f"- **{src}** -> _{rel}_ -> **{tgt}**")
        return "\n".join(lines)

    # ================================================================== #
    #  MAIN ANSWER PIPELINE                                               #
    # ================================================================== #

    async def get_response(self, query: str) -> dict:
        """
        Full RAG pipeline:
          1. Decide strategy (which engines are available)
          2. Retrieve from BM25 and/or Graph in parallel
          3. Fuse context
          4. Call Groq LLM for final answer
          5. Return structured result

        Returns dict with keys:
          answer, bm25_chunks, graph_triples, strategy
        """
        logger.info(f"Query: {query}")

        # Step 1 -- Decide strategy
        strategy = self._decide_strategy(query)
        logger.info(f"Strategy: {strategy}")

        # Step 2 -- Retrieve (run in thread to not block async loop)
        bm25_chunks: List[str] = []
        graph_triples: List[Dict] = []

        if strategy in ("both", "bm25_only"):
            bm25_chunks = await asyncio.to_thread(self._retrieve_bm25, query)

        if strategy in ("both", "graph_only"):
            graph_triples = await asyncio.to_thread(self._retrieve_graph, query)

        # Step 3 -- Format context
        bm25_ctx = self._format_bm25_context(bm25_chunks)
        graph_ctx = self._format_graph_context(graph_triples)

        combined_context = (
            "## Document Chunks (BM25 Keyword Search)\n\n"
            f"{bm25_ctx}\n\n"
            "---\n\n"
            "## Knowledge-Graph Facts (Neo4j)\n\n"
            f"{graph_ctx}"
        )

        # Step 4 -- Call LLM
        if strategy == "none" and not bm25_chunks and not graph_triples:
            answer_text = (
                "I don't have any knowledge base loaded yet. "
                "Please upload a document first using the attachment button."
            )
        else:
            messages = [
                {"role": "system", "content": SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": (
                        f"### Context\n\n{combined_context}\n\n"
                        f"---\n\n### Question\n\n{query}"
                    ),
                },
            ]
            try:
                completion = await asyncio.to_thread(
                    lambda: self.llm.chat.completions.create(
                        model=LLM_MODEL,
                        messages=messages,
                        temperature=0.4,
                        max_tokens=1024,
                        stream=False,
                    )
                )
                answer_text = completion.choices[0].message.content.strip()
            except Exception as e:
                logger.error(f"LLM generation failed: {e}")
                answer_text = (
                    "I encountered an error while generating the answer. "
                    "Please try again."
                )

        return {
            "answer": answer_text,
            "bm25_chunks": bm25_chunks,
            "graph_triples": graph_triples,
            "strategy": strategy,
        }

    # ================================================================== #
    #  FILE INGESTION (from UI uploads)                                   #
    # ================================================================== #

    async def ingest_file(self, file_name: str, file_path: str) -> dict:
        """
        Process an uploaded file:
          1. Copy raw file to data/raw/
          2. Run cleaner + chunker -> data/processed/*.chunks.jsonl
          3. Rebuild BM25 index and reload it
        """
        ext = Path(file_name).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            return {
                "status": "error",
                "message": f"Unsupported file type '{ext}'. Supported: {', '.join(SUPPORTED_EXTENSIONS)}",
                "chunks_count": 0,
            }

        try:
            dest = RAW_DIR / file_name
            shutil.copy2(file_path, dest)
            logger.info(f"Saved uploaded file to {dest}")

            chunks_count = await asyncio.to_thread(self._run_ingestion, dest)
            await asyncio.to_thread(self._rebuild_and_reload_bm25)

            return {
                "status": "ok",
                "message": f"Processed '{file_name}' -> {chunks_count} chunks. BM25 index rebuilt.",
                "chunks_count": chunks_count,
            }

        except Exception as e:
            logger.error(f"Ingestion failed for {file_name}: {e}")
            return {
                "status": "error",
                "message": f"Failed to process '{file_name}': {str(e)}",
                "chunks_count": 0,
            }

    # ================================================================== #
    #  PRIVATE HELPERS                                                    #
    # ================================================================== #

    @staticmethod
    def _run_ingestion(file_path: Path) -> int:
        """Run the cleaner + chunker on a single file. Returns chunk count."""
        cleaner_path = Path(__file__).parent.parent / "ingest" / "cleaner.py"
        spec = importlib.util.spec_from_file_location("cleaner", cleaner_path)
        cleaner = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(cleaner)

        ext = file_path.suffix.lower()

        if ext == ".pdf":
            pages = cleaner.extract_text_from_pdf(file_path)
        elif ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            paragraphs = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    paragraphs.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            pages = ["\n".join(paragraphs)]
        elif ext in (".txt", ".md"):
            pages = [file_path.read_text(encoding="utf-8", errors="ignore")]
        else:
            return 0

        cleaned_text = cleaner.clean_pages(pages, source_type="pdf")
        basename = file_path.stem
        cleaner.write_cleaned_text(PROCESSED_DIR, basename, cleaned_text)
        chunks = cleaner.chunk_text_by_sentences(cleaned_text, max_tokens=500, overlap=100)
        cleaner.write_chunks_jsonl(PROCESSED_DIR, basename, file_path, chunks)

        logger.info(f"Ingested {file_path.name}: {len(chunks)} chunks")
        return len(chunks)

    def _rebuild_and_reload_bm25(self):
        """Rebuild BM25 index and reload it into this instance."""
        self.keyword_engine.build_index()
        try:
            self.keyword_engine.load_index()
            self._bm25_ready = True
            logger.info("BM25 index rebuilt and reloaded.")
        except FileNotFoundError:
            self._bm25_ready = False
            logger.warning("BM25 rebuild produced no index.")
