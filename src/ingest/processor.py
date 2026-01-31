from __future__ import annotations

import argparse
import logging
from pathlib import Path
from typing import List
import importlib.util

from tqdm import tqdm

# logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s"
)
logger = logging.getLogger(__name__)


# -------------------- import cleaner.py dynamically --------------------

def import_cleaner_module():
    cleaner_path = Path(__file__).parent / "cleaner.py"
    spec = importlib.util.spec_from_file_location("cleaner", cleaner_path)
    cleaner = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(cleaner)
    return cleaner


# -------------------- DOCX extraction --------------------

def extract_text_from_docx(path: Path) -> List[str]:
    from docx import Document

    doc = Document(path)
    paragraphs = []

    for p in doc.paragraphs:
        if p.text and p.text.strip():
            paragraphs.append(p.text.strip())

    # extract tables as text
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    text = "\n".join(paragraphs)
    return [text]


# -------------------- file discovery --------------------

def discover_files(raw_dir: Path) -> List[Path]:
    patterns = ("*.pdf", "*.docx", "*.txt", "*.md")
    files = []
    for pat in patterns:
        files.extend(raw_dir.rglob(pat))
    return sorted(set(files))


# -------------------- single-file processor --------------------

def process_file(
    file_path: Path,
    out_dir: Path,
    cleaner,
    use_ocr: bool,
    max_tokens: int,
    overlap: int
):
    ext = file_path.suffix.lower()

    # -------- extract --------
    if ext == ".pdf":
        pages = cleaner.extract_text_from_pdf(file_path, use_ocr=use_ocr)
        source_type = "pdf"

    elif ext == ".docx":
        pages = extract_text_from_docx(file_path)
        source_type = "docx"

    elif ext in (".txt", ".md"):
        pages = [file_path.read_text(encoding="utf-8", errors="ignore")]
        source_type = "pdf"

    else:
        return {"file": str(file_path), "status": "skipped"}

    # -------- clean --------
    cleaned_text = cleaner.clean_pages(pages, source_type=source_type)

    # -------- write cleaned text --------
    out_dir.mkdir(parents=True, exist_ok=True)
    basename = file_path.stem
    cleaned_path = cleaner.write_cleaned_text(out_dir, basename, cleaned_text)

    # -------- chunk --------
    chunks = cleaner.chunk_text_by_sentences(
        cleaned_text,
        max_tokens=max_tokens,
        overlap=overlap
    )

    chunks_path = cleaner.write_chunks_jsonl(
        out_dir,
        basename,
        file_path,
        chunks
    )

    return {
        "file": str(file_path),
        "status": "ok",
        "cleaned": str(cleaned_path),
        "chunks": len(chunks),
    }


# -------------------- main --------------------

def main(
    raw_dir: str,
    out_dir: str,
    use_ocr: bool,
    max_tokens: int,
    overlap: int,
):
    raw = Path(raw_dir)
    out = Path(out_dir)

    cleaner = import_cleaner_module()

    files = discover_files(raw)
    logger.info(f"Discovered {len(files)} files in {raw}")

    results = []

    for f in tqdm(files, desc="Processing files"):
        try:
            res = process_file(
                f,
                out,
                cleaner,
                use_ocr,
                max_tokens,
                overlap
            )
            results.append(res)
        except Exception as e:
            logger.error(f"Failed processing {f}: {e}")
            results.append({"file": str(f), "status": "error", "error": str(e)})

    ok = [r for r in results if r["status"] == "ok"]
    err = [r for r in results if r["status"] == "error"]

    logger.info(f"Processing complete: {len(ok)} succeeded, {len(err)} failed")


# -------------------- entrypoint --------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser("PDF / DOCX cleaner for RAG")
    parser.add_argument("--raw-dir", default="data/raw")
    parser.add_argument("--out-dir", default="data/processed")
    parser.add_argument("--ocr", action="store_true")
    parser.add_argument("--max-tokens", type=int, default=500)
    parser.add_argument("--overlap", type=int, default=100)
    args = parser.parse_args()

    main(
        args.raw_dir,
        args.out_dir,
        args.ocr,
        args.max_tokens,
        args.overlap,
    )
